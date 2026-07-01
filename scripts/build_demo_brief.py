# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Documento-presentacion-demo-NovaDesk-AI.docx"


BLUE = "0F766E"
DARK = "111827"
MUTED = "64748B"
FILL = "F3F8F8"
HEADER_FILL = "E3F3F1"
BORDER = "D9E2E4"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for side, value in values.items():
        tag = "w:" + side
        node = margins.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=11, bold=False, color=DARK):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        style_run(r2)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    style_run(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level <= 2 else DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    style_run(r)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, FILL)
    set_cell_border(cell)
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    style_run(r, size=11, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    style_run(r2, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Demo inicial: Plataforma SaaS de automatización con IA")
    style_run(r, size=22, bold=True, color=DARK)

    add_callout(
        doc,
        "Propósito de esta versión",
        "Esta versión no tiene nada que ver con la versión real del producto. Su objetivo es únicamente dar una idea visual y funcional de lo que se quiere construir: cómo podría verse la plataforma, qué módulos tendría y cómo sería el flujo principal de trabajo. Es una demo inicial para conversar la idea, no una versión final ni una implementación productiva.",
    )

    add_heading(doc, "Link de la demo", 1)
    add_paragraph(doc, "URL actual: http://localhost:5173")

    add_heading(doc, "Qué busca mostrar la demo", 1)
    add_bullet(doc, "Cómo se vería una plataforma SaaS para centralizar mensajes de distintos canales.")
    add_bullet(doc, "Cómo un chatbot entrenable podría apoyar ventas, soporte y seguimiento.")
    add_bullet(doc, "Cómo se podrían configurar reglas de automatización y escalamiento a humanos.")
    add_bullet(doc, "Cómo el negocio podría medir volumen, rendimiento, satisfacción y ahorro de tiempo.")

    add_heading(doc, "Modulos principales", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [Inches(1.55), Inches(2.55), Inches(2.7)]
    set_table_width(table, widths)
    repeat_table_header(table.rows[0])
    headers = ["Módulo", "Qué muestra", "Para qué sirve"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(text)
        style_run(r, size=10.5, bold=True, color=BLUE)

    rows = [
        (
            "Dashboard",
            "Indicadores generales de conversaciones, mensajes automáticos, clientes atendidos, satisfacción y rendimiento del bot.",
            "Da una vista rápida del valor del sistema y permite entender el estado general de la operación.",
        ),
        (
            "Inbox",
            "Lista de conversaciones, chat central y ficha del cliente con notas, etiquetas e historial.",
            "Muestra cómo el equipo podría atender mensajes desde varios canales en un solo lugar.",
        ),
        (
            "Automatizaciones",
            "Reglas visuales para responder preguntas frecuentes, enviar catálogos, pedir datos o derivar a un agente.",
            "Explica cómo se reduciría trabajo manual y cómo se escalarían casos importantes.",
        ),
        (
            "Entrenamiento del Bot",
            "Personalidad, tono, objetivos, base de conocimiento, documentos y recomendaciones de IA.",
            "Muestra la idea de entrenar el bot según la empresa, sus productos y sus políticas.",
        ),
        (
            "Clientes",
            "Mini CRM con clientes, canal favorito, estado, interés, valor estimado y responsable.",
            "Ayuda a visualizar cómo la plataforma puede conectar atención, ventas y seguimiento comercial.",
        ),
        (
            "Canales",
            "Estado de WhatsApp, Instagram, Web Chat, Messenger y Email, con sincronización y mensajes procesados.",
            "Permite explicar que el producto busca operar como una bandeja multicanal.",
        ),
        (
            "Analiticas",
            "Gráficos de volumen, canales, horas pico, satisfacción, resolución automática y ventas asistidas.",
            "Sirve para mostrar cómo el cliente podría medir resultados y justificar el uso de la herramienta.",
        ),
        (
            "Simulador IA",
            "Chat de prueba con análisis de intención, sentimiento, confianza, entidad detectada y acción sugerida.",
            "Permite demostrar de forma simple cómo podría comportarse el bot ante distintos escenarios.",
        ),
    ]

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            style_run(r, size=10.2, bold=(idx == 0), color=DARK if idx != 0 else BLUE)

    add_heading(doc, "Ruta sugerida para presentar la demo", 1)
    steps = [
        "Abrir el Dashboard para explicar el concepto general y los indicadores.",
        "Entrar al Inbox para mostrar el flujo principal de atención.",
        "Pasar a Automatizaciones para explicar cómo se reduciría trabajo repetitivo.",
        "Mostrar Entrenamiento del Bot para explicar cómo se adapta a cada negocio.",
        "Cerrar con Simulador IA para probar un caso de venta, soporte o reclamo.",
    ]
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(step)
        style_run(r)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
